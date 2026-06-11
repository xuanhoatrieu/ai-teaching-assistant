#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Helpers for cell styling via OXML
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_left_border(cell, color_hex="3b82f6", sz="36"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>
            <w:top w:val="none"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)

def set_table_borders(table, color_hex="cbd5e1"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
            <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def add_formatted_text(paragraph, text, default_monospace=False):
    # Regex to find bold (**text**), italic (*text*), and inline code (`code`)
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            if default_monospace:
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            if default_monospace:
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(199, 37, 78) # Markdown magenta-pink color
        else:
            run = paragraph.add_run(part)
            if default_monospace:
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(71, 85, 105)

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    
    if level == 1:
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(15, 23, 42) # Slate 900
    elif level == 2:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
    elif level == 3:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(51, 65, 85) # Slate 700
    elif level == 4:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(71, 85, 105) # Slate 600

def add_code_block(doc, code_text, lang=""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    
    set_cell_background(cell, "f8fafc") # Slate 50
    
    # Custom thin border
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="e2e8f0"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="e2e8f0"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="e2e8f0"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="e2e8f0"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    cell_p = cell.paragraphs[0]
    cell_p.paragraph_format.space_before = Pt(4)
    cell_p.paragraph_format.space_after = Pt(4)
    cell_p.paragraph_format.line_spacing = 1.0
    
    code_lines = code_text.split('\n')
    for idx, line in enumerate(code_lines):
        if idx > 0:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
        else:
            p = cell_p
        
        add_formatted_text(p, line, default_monospace=True)

def add_callout(doc, callout_lines):
    callout_type = "NOTE"
    color_hex = "3b82f6" # default blue
    bg_hex = "eff6ff"
    
    first_line = callout_lines[0] if callout_lines else ""
    type_match = re.match(r'^\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\](.*)$', first_line, re.IGNORECASE)
    
    if type_match:
        callout_type = type_match.group(1).upper()
        first_line_rest = type_match.group(2).strip()
        callout_lines[0] = first_line_rest
        
        if callout_type == "IMPORTANT":
            color_hex = "6366f1" # indigo
            bg_hex = "f5f3ff"
        elif callout_type == "WARNING":
            color_hex = "f59e0b" # amber
            bg_hex = "fffbeb"
        elif callout_type == "CAUTION":
            color_hex = "f43f5e" # rose
            bg_hex = "fff1f2"
        elif callout_type == "TIP":
            color_hex = "10b981" # emerald
            bg_hex = "f0fdf4"
            
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    
    set_cell_background(cell, bg_hex)
    set_cell_left_border(cell, color_hex, sz="36")
    
    cell_p = cell.paragraphs[0]
    cell_p.paragraph_format.space_before = Pt(4)
    cell_p.paragraph_format.space_after = Pt(2)
    
    # Add title run
    r_color = RGBColor(int(color_hex[0:2],16), int(color_hex[2:4],16), int(color_hex[4:6],16))
    run_title = cell_p.add_run(f"✦ {callout_type.capitalize()}: ")
    run_title.bold = True
    run_title.font.color.rgb = r_color
    
    if callout_lines[0]:
        add_formatted_text(cell_p, callout_lines[0])
        start_idx = 1
    else:
        start_idx = 1 if len(callout_lines) > 1 else 0
        if len(callout_lines) == 1:
            start_idx = 0
            
    for line in callout_lines[start_idx:]:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        add_formatted_text(p, line)

def add_table(doc, table_lines):
    rows_data = []
    for line in table_lines:
        parts = [p.strip() for p in line.split('|')]
        if parts and parts[0] == '':
            parts = parts[1:]
        if parts and parts[-1] == '':
            parts = parts[:-1]
        rows_data.append(parts)
        
    if not rows_data:
        return
        
    if len(rows_data) > 1 and all(re.match(r'^:?-+:?$', cell) for cell in rows_data[1]):
        rows_data.pop(1) # Remove alignment line
        
    headers = rows_data[0]
    data_rows = rows_data[1:]
    num_cols = len(headers)
    
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    hdr_cells = table.rows[0].cells
    for col_idx, text in enumerate(headers):
        hdr_cells[col_idx].text = ""
        p = hdr_cells[col_idx].paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[col_idx], "1e293b") # Slate-800
        
    for r_idx, row_parts in enumerate(data_rows):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "f8fafc" if r_idx % 2 == 0 else "ffffff" # Slate-50 zebra stripe
        for c_idx in range(num_cols):
            val = row_parts[c_idx] if c_idx < len(row_parts) else ""
            row_cells[c_idx].text = ""
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_text(p, val)
            if bg_color != "ffffff":
                set_cell_background(row_cells[c_idx], bg_color)

def main():
    if len(sys.argv) < 2:
        print("Usage: python3 md_to_docx.py <input.md> [output.docx]")
        sys.exit(1)
        
    input_file = sys.argv[1]
    if len(sys.argv) > 2:
        output_file = sys.argv[2]
    else:
        output_file = os.path.splitext(input_file)[0] + ".docx"
        
    if not os.path.exists(input_file):
        print(f"Error: {input_file} does not exist!")
        sys.exit(1)
        
    print(f"Reading {input_file}...")
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()
        
    doc = Document()
    
    # Standard margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles config
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(51, 65, 85) # Slate-700
    
    lines = content.split('\n')
    i = 0
    total_lines = len(lines)
    
    while i < total_lines:
        line = lines[i]
        stripped = line.strip()
        
        # Code blocks
        if stripped.startswith('```'):
            code_lines = []
            lang = stripped[3:].strip()
            i += 1
            while i < total_lines and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1 # skip ending ```
            add_code_block(doc, '\n'.join(code_lines), lang)
            continue
            
        # Tables
        if stripped.startswith('|'):
            table_lines = []
            while i < total_lines and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, table_lines)
            continue
            
        # Callouts
        if stripped.startswith('>'):
            callout_lines = []
            while i < total_lines and (lines[i].strip().startswith('>') or (lines[i].strip() == '' and callout_lines)):
                if lines[i].strip() == '':
                    if i + 1 < total_lines and lines[i+1].strip().startswith('>'):
                        callout_lines.append('')
                        i += 1
                        continue
                    else:
                        break
                callout_lines.append(lines[i].strip().lstrip('>').strip())
                i += 1
            add_callout(doc, callout_lines)
            continue
            
        # Headings
        if stripped.startswith('# '):
            add_heading(doc, stripped[2:], 1)
            i += 1
            continue
        elif stripped.startswith('## '):
            add_heading(doc, stripped[3:], 2)
            i += 1
            continue
        elif stripped.startswith('### '):
            add_heading(doc, stripped[4:], 3)
            i += 1
            continue
        elif stripped.startswith('#### '):
            add_heading(doc, stripped[5:], 4)
            i += 1
            continue
            
        # Bullet list items
        bullet_match = re.match(r'^(\s*)([-*]|\+)\s+(.*)$', line)
        if bullet_match:
            indent_spaces = len(bullet_match.group(1))
            text_val = bullet_match.group(3)
            level = 0
            if indent_spaces >= 3:
                level = 1
            if indent_spaces >= 6:
                level = 2
            
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_text(p, text_val)
            i += 1
            continue
            
        # Numbered list items
        num_match = re.match(r'^(\s*)(\d+)\.\s+(.*)$', line)
        if num_match:
            indent_spaces = len(num_match.group(1))
            text_val = num_match.group(3)
            level = 0
            if indent_spaces >= 3:
                level = 1
                
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_text(p, text_val)
            i += 1
            continue
            
        # Horizontal rules
        if stripped in ('---', '***', '___'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("―" * 40)
            run.font.color.rgb = RGBColor(203, 213, 225) # Slate-300
            i += 1
            continue
            
        if stripped == '':
            i += 1
            continue
            
        # Standard Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        add_formatted_text(p, stripped)
        i += 1

    print(f"Saving to {output_file}...")
    doc.save(output_file)
    print("Success!")

if __name__ == '__main__':
    main()
